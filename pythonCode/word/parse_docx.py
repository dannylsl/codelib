import os
import sys
import win32com
#from win32com.client import Dispatch, constants
from docx.api import Document
import MySQLdb
import datetime


if __name__ == "__main__" :

#    w = win32com.client.Dispatch('Word.Application')
    print "main"
    filename = "test.docx"
    path = "C:\\Users\\dannylsl\\Documents\\"
    filepath = "%s%s"%(path, filename)
    print filepath

    #docx = w.Documents.Open( FileName = f )
    d = Document(filepath)
    print "TABLE COUNTS:",len(d.tables)
    table_index = len(d.tables)/2
    for i in range(0, table_index) :
        r_len_h = len(d.tables[i].rows)
        c_len_h = len(d.tables[i].columns)

        r_len_b = len(d.tables[i+1].rows)
        c_len_b = len(d.tables[i+1].columns)
        print "Head rows=%s  columns=%s"%(r_len_h, c_len_h)
        print "Body rows=%s  columns=%s"%(r_len_h, c_len_h)

        #TABLE Header
        if r_len_h != 4 and c_len_h != 2 :
            print "Head error"
            print "[%3d] rows=%3d  columns=%3d"%(i ,r_len_h, c_len_h)
            sys.exit()

        name    = d.tables[i].cell(0,1).text
        code    = d.tables[i].cell(1,1).text
        comment = d.tables[i].cell(2,1).text
        primary = d.tables[i].cell(3,1).text
        print name,code,comment,primary

        #TABLE Body
        if c_len_b != 4 :
            print "Body error"
            print "[%3d] rows=%3d  columns=%3d"%(i ,r_len_b, c_len_b)
            sys.exit()

        for r in range(1, r_len_b) :
            for c in range(0, c_len_b) :
                print "[%s, %s]%s"%(r,c, d.tables[i+1].cell(r,c).text)

    sys.exit()
